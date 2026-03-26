# your_app/utils/randomize.py
import os, shutil, zipfile, random, string, re, uuid, base64, tempfile
from datetime import datetime
from threading import Timer
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement, qn
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.shared import Inches
# from docx2pdf import convert
from django.conf import settings
from .timestamp import get_timestamp
# from .check_time import delete_time
from hooks.pretty_print import pretty_print_json
from latex2mathml.converter import convert
from lxml import etree
from school.models import ScrambleSession

# # Ensure /public folder exists
# PUBLIC_DIR = os.path.join(settings.BASE_DIR, 'public')
# os.makedirs(PUBLIC_DIR, exist_ok=True)

get_abbr_object = {
	"subject_abbr": {
		"english-language": "ENG",
		"mathematics": "MTH",
		"civic-education": "CVE",
		"basic-science": "BSC",
		"social-studies": "SOS",
		"crs": "CRS",
		"irs": "IRS",
		"history": "HIS",
		"cultural-&-creative-arts": "CCA",
		"computer-studies": "CST",
		"home-economics": "HME",
		"one-nigerian-language": "ONL",
		"integrated-science": "INS",
		"physical-&-health-education": "PHE",
		"digital-technologies": "DGT",
		"crs-/-irs": "CRI",
		"nigerian-history": "NHS",
		"social-&-citizenship-studies": "SCS",
		"business-studies": "BST",
		"french-studies": "FRN",
		"arabic-studies": "ARB",
		"biology": "BIO",
		"chemistry": "CHM",
		"physics": "PHY",
		"further-mathematics": "FMTH",
		"agricultural-science": "AGR",
		"technical-drawing": "TD",
		"geography": "GEO",
		"computer-studies-/-ict": "ICT",
		"physical-/-health-education": "PHE",
		"foods-&-nutrition-/-home-economics": "FNH",
		"literature-in-english": "LIT",
		"government": "GOV",
		"french-/-other-foreign-languages": "FOL",
		"nigerian-language(s)": "NLA",
		"visual-/-fine-arts": "VFA",
		"music": "MUS",
		"economics": "ECO",
		"commerce": "COM",
		"financial-accounting": "FAC",
		"marketing": "MKT",
		"accounting": "ACC",
		"office-practice": "OFP",
		"bookkeeping": "BKP",
		"data-processing-/-computer-studies": "DPC",
	},
	"term_abbr": {
		"first": "1st",
		"second": "2nd",
		"third": "3rd",
	}
}
def get_abbr(category, value):
	print(f'value: {value}')
	if not value:
		return value
	return get_abbr_object[category].get(
		value.strip().lower(),
		value
	)

def get_shuffle_record(db_category):
	print('checking ')
	if db_category:
		print('checking if shuffle_record exists')
		queryset_record = ScrambleSession.objects.filter(**db_category)
		print(f'queryset_record: {queryset_record}')
		record = queryset_record.values_list('shuffle_record', flat=True).first()
		print(f'record:')
		pretty_print_json(record)
		return record, False if queryset_record else True
	else:
		print('no category data provided')
		return None

def base64_to_png_file(base64_str):
	"""
	Converts a base64 image string (data:image/png;base64,...) into a temporary PNG file.
	Returns the file path.
	"""
	if base64_str.startswith("data:image/png;base64,"):
		base64_str = base64_str.split(",")[1]  # strip header

	image_bytes = base64.b64decode(base64_str)
	tmp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
	tmp_file.write(image_bytes)
	tmp_file.close()
	return tmp_file.name

latex_pattern = r'(\\[a-zA-Z]+.*)'
def mathml_to_omml(mathml_string):
	xslt_path = os.path.join(settings.BASE_DIR, "MML2OMML.XSL")
	# print(f'xslt_path: {xslt_path}')
	# print(f'path exist: {os.path.exists(xslt_path)}')

	mathml_dom = etree.fromstring(mathml_string.encode())

	# print(f'mathml_dom: {mathml_dom}')
	xslt = etree.parse(xslt_path)
	# print(f'xslt: {xslt}')
	transform = etree.XSLT(xslt)
	# print(f'transform: {transform}')

	omml_dom = transform(mathml_dom)
	# print(f'omml_dom: {omml_dom}')
	return omml_dom.getroot()

def shuffle_array(arr, return_order=False, order=None):
	if isinstance(arr, dict):
		arr = list(arr.values())
	else:
		arr = arr[:]

	# If we already have an order, reuse it
	if order:
		shuffled = [arr[i] for i in order]
		return shuffled, None

	# Otherwise generate a new shuffle
	indices = list(range(len(arr)))
	random.shuffle(indices)

	shuffled = [arr[i] for i in indices]

	if return_order:
		return shuffled, indices

	return shuffled, None

def generate_alphabets(n):
	return [chr(65 + i) for i in range(n)]  # A, B, C, D...

def save_docx(
	paragraphs,
	file_path,
	add_footer=True,
	image_map=None,
	logo=None,
	extracts=None):
	print('✅✅✅✅✅✅✅✅✅✅')
	print(f'paragraphs:')
	pretty_print_json(paragraphs)
	# print(f'extracts:')
	# pretty_print_json(extracts)
	print(f'file_path: {file_path}')
	print(f'add_footer: {add_footer}')
	print(f'image_map: {image_map}')
	print(f'logo: {logo}')
	print('❎❎❎❎❎❎❎❎❎❎')
	# #####################
	# print("data to use for file:")
	# pretty_print_json(paragraphs)
	# #####################
	if image_map is None:
		image_map = {}
	doc = Document()
	# ============================
	# PAGE SIZE & MARGINS (PUT IT HERE)
	# ============================
	from docx.shared import Inches

	for section in doc.sections:
		# section.page_width = Inches(11.69)   # A4 landscape
		# section.page_height = Inches(8.27)

		section.left_margin = Inches(0.7)
		section.right_margin = Inches(0.7)
		section.top_margin = Inches(0)
		section.bottom_margin = Inches(0.75)

	# ============================
	# SECTION 1 — PAGE 1 HEADER (applicable to only the first page)
	# ============================
	section1 = doc.sections[0]
	section1.different_first_page_header_footer = True

	header = section1.first_page_header

	# ============================
	# HEADER PARAGRAPH WITH RHS LOGO
	# ============================
	p = header.paragraphs[0]
	p.clear()

	# right-aligned tab stop at page margin
	pPr = p._p.get_or_add_pPr()
	tabs = OxmlElement('w:tabs')
	tab = OxmlElement('w:tab')
	tab.set(qn('w:val'), 'right')
	tab.set(qn('w:pos'), str(int(Inches(7.2).emu)))  # usable page width
	tabs.append(tab)
	pPr.append(tabs)

	# column widths
	# left_cell.width = Inches(5.5)
	# right_cell.width = Inches(1.5)

	# ============================
	# LOGO (HEADER - FIRST PAGE ONLY)
	# ============================
	if logo:
		logo_para = header.add_paragraph()
		logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

		run = logo_para.add_run()
		run.add_picture(
			logo,
			width=Inches(0.7)  # adjust if needed
		)

		logo_para.paragraph_format.space_before = Pt(0)
		logo_para.paragraph_format.space_after = Pt(1)
		# run = p.add_run("\t")
		# run.add_picture(
		# 	logo,
		# 	width=Inches(1.2)  # logo size ONLY
		# )
	# header_para = header.paragraphs[0]

	header_lines = []
	body_start_index = 0

	for i, line in enumerate(paragraphs):
		if line == "":
			body_start_index = i + 1
			break
		header_lines.append(line)

	# Clear default paragraph
	# header.paragraphs[0].clear()
	# header.paragraphs[0].clear()

	school = None
	if len(header_lines)==8:
		school, subject, clazz, term, duration, instruction, variant, qtype = header_lines
	else:
		subject, clazz, term, duration, instruction, variant, qtype = header_lines

	# --- School name (centered & bold)
	if school:
		p = header.add_paragraph(school)
		p.alignment = WD_ALIGN_PARAGRAPH.CENTER
		p.runs[0].bold = True
		p.runs[0].font.size = Pt(14)
		p.paragraph_format.space_before = Pt(0)
		p.paragraph_format.space_after = Pt(2)
		p.paragraph_format.line_spacing = 0.5
		p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

	# --- Subject (left)
	p = header.add_paragraph()
	p.add_run(subject).bold = True
	p.alignment = WD_ALIGN_PARAGRAPH.LEFT
	p.paragraph_format.space_before = Pt(0)
	p.paragraph_format.space_after = Pt(2)
	p.paragraph_format.line_spacing = 0.5
	p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

	# --- Term (left) | Class (right)
	p = header.add_paragraph()
	p.add_run(term)
	p.add_run("\t" * 8)
	p.add_run(clazz)
	p.alignment = WD_ALIGN_PARAGRAPH.LEFT
	p.paragraph_format.space_before = Pt(0)
	p.paragraph_format.space_after = Pt(2)
	p.paragraph_format.line_spacing = 0.5
	p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

	# --- Variant ID (left)
	p = header.add_paragraph()
	p.add_run(variant)
	p.paragraph_format.space_before = Pt(0)
	p.paragraph_format.space_after = Pt(2)
	p.paragraph_format.line_spacing = 0.5
	p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

	# --- Duration (left) | Type (center) | Instruction (right)
	p = header.add_paragraph()
	p.add_run(duration)
	p.add_run("\t" * 2)
	p.add_run(qtype)
	p.alignment = WD_ALIGN_PARAGRAPH.LEFT
	p.paragraph_format.space_before = Pt(0)
	p.paragraph_format.space_after = Pt(2)
	p.paragraph_format.line_spacing = 0.5
	p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
	p.add_run("\t" * 1)
	p.add_run(instruction).bold = True
	p.alignment = WD_ALIGN_PARAGRAPH.LEFT
	p.paragraph_format.space_before = Pt(0)
	p.paragraph_format.space_after = Pt(2)
	p.paragraph_format.line_spacing = 0.5
	p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

	# ============================
	# SECTION 2 — QUESTIONS (2 COLUMNS)
	# ============================
	section2 = doc.add_section(WD_SECTION.CONTINUOUS)
	sectPr = section2._sectPr

	cols = OxmlElement('w:cols')
	cols.set(qn('w:num'), '2')
	cols.set(qn('w:space'), '720')
	sectPr.append(cols)

	# --- Add question body content (flows across columns & pages)
	for para in paragraphs[body_start_index:]:
		# print(f'para: {para}')
		# Detect image marker
		if para.startswith("__IMAGE__:"):
			print(f'__IMAGE__')
			q_index = int(para.split(":")[1])
			image = image_map.get(q_index)

			if image:
				img_paragraph = doc.add_paragraph()
				run = img_paragraph.add_run()

				# Add image (COLUMN-FRIENDLY SIZE)
				run.add_picture(
					image,
					width=Inches(2.5)  # fits 2-column layout
				)

				img_paragraph.paragraph_format.space_after = Pt(6)
				img_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

			continue  # skip text rendering for marker

		# print(f'para: {para}')

		mathml = None
		match = re.search(latex_pattern, para)
		if match:
			latex = match.group(1)
			normal_text = para.replace(latex, '').strip()

			# --- TEXT LINE ---
			if normal_text:
				p_text = doc.add_paragraph(normal_text)
				p_text.runs[0].font.size = Pt(12)

			# --- MATH LINE (separate paragraph) ---
			math_paragraph = doc.add_paragraph()

			mathml = convert(latex)
			# print(f'mathml: {mathml}')
			omml = mathml_to_omml(mathml)
			# print(f'omml: {omml}')

			math_paragraph._p.append(omml)
			# run = math_paragraph.add_run()   # create a run
			# if omml.tag.endswith("oMathPara"):
			# 	omath = omml.find(".//{http://schemas.openxmlformats.org/officeDocument/2006/math}oMath")
			# 	if omath is not None:
			# 		run._r.append(omath)
			# else:
			# 	run._r.append(omml)

		else:
			p = doc.add_paragraph(para)
			# p.runs[0].font.size = Pt(12)
		# normal text
		# p = doc.add_paragraph(para)
		# p = doc.add_paragraph()
		# if normal_text:
		# 	p.add_run(normal_text + " ")
		# if latex:
		# 	mathml = convert(latex)
		# 	omml = mathml_to_omml(mathml)

		# 	p._p.append(omml)
		p.style.font.size = Pt(12)

		# # if math
		# # normal text
		# p = doc.add_paragraph(para)
		# p.style.font.size = Pt(12)

		# Detect options: A. B. C. D.
		if re.match(r'^[A-D]\.\s', para):
			p.paragraph_format.space_before = Pt(0)
			p.paragraph_format.space_after = Pt(2)
			p.paragraph_format.line_spacing = 0.5
			p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

	# ============================
	# SECTION 3 — LAST PAGE FOOTER (OPTIONAL)
	# ============================
	if add_footer:
		section3 = doc.add_section()
		section3.start_type = 2  # new page

		footer = section3.footer
		footer_para = footer.paragraphs[0]
		footer_para.text = "End of Examination"
		footer_para.style.font.size = Pt(10)
		footer_para.alignment = 1  # center

	doc.save(file_path)

def get_unique_id():
	random_str = ''.join(random.choices(string.ascii_lowercase + string.digits, k=4))
	print(f"Generated random string for variant ID: {random_str}")
	# pretty_print_json(random_str)

	# Get current time
	now = datetime.now()
	year = now.year
	month = f"{now.month:02d}"
	date = f"{now.day:02d}"
	hours = f"{now.hour:02d}"
	minutes = f"{now.minute:02d}"
	seconds = f"{now.second:02d}"
	return f"{year}{month}{date}_{hours}{minutes}{seconds}_{random_str}"

term_map = {
	"first": "1st",
	"second": "2nd",
	"third": "3rd"
}
def clean_term(term):
	if not term:
		return "no_term-term"
	term_assignment = term_map.get(term.lower(), "no_term")
	return f"{term_assignment}-term"

def Randomize(data, multiple=False, db_category=None):
	try:
		print('db_category:')
		pretty_print_json(db_category)
		saved_shuffle_record, gen_new = get_shuffle_record(db_category)
		print('saved_shuffle_record:')
		print(saved_shuffle_record)
		# random_str = ''.join(random.choices(string.ascii_lowercase + string.digits, k=4))
		# variant_id = f"{uuid.uuid4().hex[:8]}_{random_str}"
		# Generate 4 random lowercase alphanumeric characters
		
		# get subject
		subject = data['subject'].lower()
		generated_term = clean_term(data.get('term', None))

		# Combine to form variantId
		unique_time = get_unique_id()
		variant_id = f"{subject}_{generated_term}_{unique_time}"
		subject_abbr = get_abbr("subject_abbr", subject)
		term_abbr = get_abbr("term_abbr", generated_term.split("-")[0])
		_day, _time, _str = unique_time.split("_")
		exam_id = f"{subject_abbr}_{term_abbr}_{_time}_{_str}"
		print(f'Generated Exam ID: {exam_id}')
		print(f"Generated variant ID: {variant_id}")

		# Create directories and files names
		zip_filename = f"estudiebuddie_{variant_id}.zip"
		print("ZIP filename will be:")
		pretty_print_json(zip_filename)
		public_dir = os.path.join(settings.BASE_DIR, 'public')
		dir_path = os.path.join(public_dir, variant_id)
		os.makedirs(dir_path, exist_ok=True)
		os.chmod(public_dir, 0o777)
		os.chmod(dir_path, 0o777)

		# Generate question types (A-Z)
		types = generate_alphabets(int(data['noOfTypes']))
		print(f"Generated question types:")
		pretty_print_json(types)

		# compare the length of incoming questions to the saved
		saved_questions_length = len(saved_shuffle_record["A"]["question_order"]) if (saved_shuffle_record and not gen_new) else 0
		print(f'saved_questions_length: {saved_questions_length}')
		incoming_questions_length = len(data.get('postQuestions', data.get('questions')).values())
		print(f'incoming_questions_length: {incoming_questions_length}')
		reshuffle = saved_questions_length != incoming_questions_length
		# Generate question and answer files for each type
		all_shuffle_records = {}
		shuffle_record = None
		for i, type_code in enumerate(types):
			print(f"Generating files for question type: {type_code}")
			# print(f"data: {data}")
			# print(f"postQuestions: {data['postQuestions']}")
			# Shuffle questions
			# questions = shuffle_array(data.get('postQuestions', data.get('questions')))
			order = None
			if not reshuffle and saved_shuffle_record and not gen_new:
				order = saved_shuffle_record.get(type_code, {}).get("question_order")
			questions, question_order = shuffle_array(
				data.get('postQuestions', data.get('questions')),
				return_order=True,
				order=order
			)
			if question_order:
				shuffle_record = {
					"question_order": question_order,
					"option_orders": {}
				}
			answer_key = [
				f"variant: {exam_id}",
				f"Type: {type_code}",
				""
			]
			print(f"Initialized answer key for type {type_code}")
			pretty_print_json(answer_key)
			print(f"Shuffled questions for type {type_code}")
			# print(f"Shuffled Questions:")
			# pretty_print_json(questions)

			# Prepare header here
			duration = data['duration']
			try:
				# print(f"Processing duration: {duration}")
				num = float(duration)
				# print(f"Converted duration to float: {num}")
				if num > 1:
					# print(f"Duration is plural hours: {num}")
					# plural: keep as hours
					duration_str = f"{num:g} hours"   # :g removes trailing .0 for whole numbers
				elif num == 1:
					# print(f"Duration is singular one hour")
					# singular
					duration_str = f"{num:g} hour"
				elif 0 < num < 1:
					# print(f"Duration is fractional hours: {num}")
					# convert fractional hours to minutes
					minutes = round(num * 60)
					# print(f"Converted fractional hours to minutes: {minutes}")
					duration_str = f"{minutes} minute{'s' if minutes != 1 else ''}"
				else:
					# print(f"Duration is zero or negative: {num}")
					# for 0 or negative values, just leave as-is or customize
					duration_str = str(duration)
			except (ValueError, TypeError, OverflowError):
				# print(f"Duration is non-numeric: {duration}")
				duration_str = duration
			school_name = data.get("school", None)
			header_lines = [
				# data['school'].upper(),
				f"Subject: {data['subject'].title()}",
				f"Class: {data['class'].upper()}",
				f"{data['term'].title()} Term",
				f"Duration: {duration_str}",
				f"Instruction: {data['instruction'].title()}",
				f"variant: {exam_id}",
				f"Type: {type_code}",
				""
			]
			if school_name:
				header_lines.insert(0, school_name.upper())
			# print(f"Prepared header for type {type_code}")
			# print(f"Header Lines:")
			# pretty_print_json(header_lines)

			# Prepare questions with shuffled options
			question_lines = header_lines[:]
			# print(f"Starting to process questions for type {type_code}")
			# print(f"Initial Question Lines:")
			# # pretty_print_json(question_lines)

			extracted_full_questions_with_types = None

			# Process each question
			for idx, q in enumerate(questions):
				# print(f"Processing question {idx + 1} for type {type_code}")
				# print(f"Question Data:")
				# pretty_print_json(q)
				question_obj = [qo for qo in q if qo]
				q = question_obj[0]
				# print('extrated question:')
				# pretty_print_json(q)
				# Shuffle options
				# opts = shuffle_array([
				# 	{"text": q["correct_answer"], "isCorrect": True},
				# 	{"text": q["wrong_answer1"], "isCorrect": False},
				# 	{"text": q["wrong_answer2"], "isCorrect": False},
				# 	{"text": q["wrong_answer3"], "isCorrect": False},
				# ])
				if not reshuffle and saved_shuffle_record and not gen_new:
					order = saved_shuffle_record.get(type_code, {}).get("option_orders", {}).get(str(idx))
				opts, option_order = shuffle_array([
					{"text": q["correct_answer"], "isCorrect": True},
					{"text": q["wrong_answer1"], "isCorrect": False},
					{"text": q["wrong_answer2"], "isCorrect": False},
					{"text": q["wrong_answer3"], "isCorrect": False},
					], return_order=True,
					order=order
				)
				if option_order:
					shuffle_record["option_orders"][str(idx)] = option_order

				# print(f"Shuffled options for question {idx + 1}:")
				# pretty_print_json(opts)
				# Label options A, B, C, D
				for j, opt in enumerate(opts):
					# print(f"Labeling option {j + 1} for question {idx + 1}")
					# print(f"Option Data Before Labeling:")
					# pretty_print_json(opt)
					opt['label'] = chr(65 + j)
					# print(f"Option Data After Labeling:")
					# pretty_print_json(opt)

				# Find correct option for answer key
				correct = next(o for o in opts if o["isCorrect"])
				# print(f"Correct option for question {idx + 1}:")
				# pretty_print_json(correct)
				# Append to answer key
				answer_key.append(f"{idx + 1}. {correct['label']}")
				# print(f"Updated answer key:")
				# pretty_print_json(answer_key)

				# append image (if it exists)
				if q.get("image"):
					question_lines.append(f"__IMAGE__:{idx}")
				# Append question and options to question lines
				# print('question:')
				# pretty_print_json(q)
				# extracted_question = q["question"][0][0]["value"]
				extracted_full_questions_with_types = q.get("question", None)

				result = {}
				for items in extracted_full_questions_with_types.values():
					for item in items:
						if item and item.get("type") and item.get("value"):
							result[item["type"]] = item["value"]

				extracted_full_questions_with_types = result
				# print('🏆🏆🏆🏆🏆🏆🏆🏆🏆+++++')
				# print(f'extracted_full_questions_with_types:')
				# pretty_print_json(extracted_full_questions_with_types)
				# print('🏆🏆🏆🏆🏆🏆🏆🏆🏆-----')

				diagram_line = extracted_full_questions_with_types.get("diagram", None)
				if diagram_line:
					print('appending diagram:')
					question_lines.append(f"__IMAGE__:{idx}")

				# appending current question
				question_lines.append(f"{idx + 1}. {extracted_full_questions_with_types['text']}")

				math_line = extracted_full_questions_with_types.get("math", None)
				if math_line:
					if math_line.startswith("{") and math_line.endswith("}"):
						print(f'original math line: {math_line}')
						math_line = math_line[1:-1]
						print(f'no_curly: {math_line}')

					print(f'adding math line: {math_line}')
					question_lines.append(f"{''.rjust(2, ' ')} {math_line}")
				print(f"Added question {idx + 1} to question lines.")
				# print('🎲🎲🎲🎲🎲🎲🎲')
				# print(f"Current Question Lines:")
				# pretty_print_json(question_lines)
				# print('🥏🥏🥏🥏🥏🥏🥏')

				# Loop and append options
				for opt in opts:
					# print(f"Adding option for question {idx + 1}")
					# pretty_print_json(opt)
					# append option label and text
					question_lines.append(f"{opt['label']}. {opt['text']}")
					# print(f"Added option`s label to question lines.")
					# pretty_print_json(opt["label"])
				# Add a blank line after each question
				question_lines.append("")
				print(f"Finished processing question {idx + 1}. Current Question Lines:")
				# pretty_print_json(question_lines)

				# stored shuffle records
				if shuffle_record:
					all_shuffle_records[type_code] = shuffle_record

			# Save Question and Answer Files
			quest_dir = os.path.join(dir_path, 'questions')
			print(f"Creating directories for type {type_code}")
			ans_dir = os.path.join(dir_path, 'answers')
			print(f"Creating answer directories for type {type_code}")
			os.makedirs(quest_dir, exist_ok=True)
			os.makedirs(ans_dir, exist_ok=True)
			question_path_for_docx = os.path.join(quest_dir, f"Question_type_{type_code}.docx")
			# question_path_for_pdf = os.path.join(quest_dir, f"Exam_type_{type_code}.pdf")

			# build image map
			image_map = {}
			for i, q in enumerate(questions):
				# print("q question:")
				# pretty_print_json(q)
				question_obj = [qo for qo in q if qo]
				extracted_question_with_image = question_obj[0]
				# print('✨✨✨✨✨✨✨✨✨000')
				
				img = None
				if extracted_question_with_image.get("image"):
					print('♻♻♻♻♻♻♻♻♻♻')
					img = extracted_question_with_image["image"]
				else:
					diagram_info = extracted_question_with_image.get("question", None)
					if diagram_info:
						diagram_info = [item for sublist in diagram_info.values() for item in sublist if item]
						
						# print('✨✨✨✨✨✨✨✨✨111')
						# print(f'diagram_info:')
						# # pretty_print_json(diagram_info)
						# print('✨✨✨✨✨✨✨✨✨---')
						dia_keys = set([k["type"] for k in diagram_info])
						# pretty_print_json(dia_keys)
						# print('diagram_info22222:')
						# pretty_print_json(diagram_info)
						if 'diagram' in dia_keys:
							diagram_in_q = [d['value'] for d in diagram_info if d['type']=='diagram_png'][0]
							# print('diagram_in_q:')
							# pretty_print_json(diagram_in_q)
							img = base64_to_png_file(diagram_in_q)
							print('img:')
							pretty_print_json(img)
							# if q.get("image"):
							# question_lines.append(f"__IMAGE__:{idx}")
					# print('✨✨✨✨✨✨✨✨✨222')

				if img:
					print('🥇🥇🥇🥇🥇🥇🥇🥇🥇🥇🥇')
					image_map[i] = img
				# if extracted_question_with_image.get("image"):
				# 	image_map[i] = extracted_question_with_image["image"]


			# print('extracted_full_questions_with_types:')
			# pretty_print_json(extracted_full_questions_with_types)
			save_docx(
				question_lines,
				question_path_for_docx,
				add_footer=False,
				image_map=image_map,
				logo=data.get("logo", None),
				extracts=extracted_full_questions_with_types,
			)

			# # pdf creation
			# try:
			# 	convert(question_path_for_docx, question_path_for_pdf)
			# 	print(f"PDF generated successfully: {question_path_for_pdf}")
			# except Exception as e:
			# 	print(f"PDF generation failed for {question_path_for_docx}: {e}")

			# with open(os.path.join(quest_dir, f"Exam_type_{type_code}.txt"), "w") as f:
			# 	print(f"Saving question text file for type {type_code}")
			# 	f.write("\n".join(question_lines))

			with open(os.path.join(ans_dir, f"Answers_type_{type_code}.txt"), "w") as f:
				print(f"Saving answer key text file for type {type_code}")
				f.write("\n".join(answer_key))
		print("FULL SHUFFLE RECORD")
		pretty_print_json(all_shuffle_records)
		if not all_shuffle_records:
			all_shuffle_records = None

		# Zip the directory if not batching
		if not multiple:
			# zip_path = os.path.join(public_dir, zip_filename)
			# print(f"Creating ZIP archive at: {zip_path}")
			# with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as archive:
			# 	print(f"Adding files to ZIP archive: {zip_path}")
			# 	for foldername, _, filenames in os.walk(dir_path):
			# 		print(f"Walking through folder: {foldername}")
			# 		for filename in filenames:
			# 			print(f"Adding file to archive: {filename}")
			# 			file_path = os.path.join(foldername, filename)
			# 			arcname = os.path.relpath(file_path, os.path.join(dir_path, '..'))
			# 			archive.write(file_path, arcname)

			# # schedule_deletion(zip_path, dir_path)

			# return f"/public/{zip_filename}"
			print('NOT MULTIPLE')
			return zip_all(dir_paths=[dir_path], zip_name=zip_filename), all_shuffle_records

		# if batch
		print('YEAH, MULTIPLE')
		return {
			"variant_id": variant_id,
			"dir_path": dir_path,
		}, all_shuffle_records

	except Exception as e:
		print(f"Error generating exam bundle: {get_timestamp()} =>", e)
		raise Exception("Failed to generate exam bundle")


def zip_all(dir_paths, zip_name=None):
	if not zip_name:
		zip_name = f"exam_bundle_multiple_{get_unique_id()}.zip"

	public_dir = os.path.join(settings.BASE_DIR, "public")
	zip_path = os.path.join(public_dir, zip_name)

	with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as archive:
		for dir_path in dir_paths:
			for foldername, _, filenames in os.walk(dir_path):
				for filename in filenames:
					file_path = os.path.join(foldername, filename)

					# keep folder structure inside zip
					arcname = os.path.relpath(
						file_path,
						os.path.dirname(dir_path)
					)
					archive.write(file_path, arcname)

	# CLEANUP (after zip is closed)
	for dir_path in dir_paths:
		if os.path.exists(dir_path):
			print(f'deleting: {dir_path}')
			shutil.rmtree(dir_path)

	return f"/public/{zip_name}"